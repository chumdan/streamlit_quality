from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import os

def convert_md_to_docx(md_file, docx_file):
    # 마크다운 파일 읽기
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # HTML로 변환
    html = markdown.markdown(md_content)
    
    # 워드 문서 생성
    doc = Document()
    
    # 제목 스타일 설정
    title = doc.add_heading('품질관리시스템 기능 설명서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 마크다운 내용을 워드로 변환
    paragraphs = html.split('\n')
    for p in paragraphs:
        if p.startswith('<h1>'):
            doc.add_heading(p[4:-5], level=1)
        elif p.startswith('<h2>'):
            doc.add_heading(p[4:-5], level=2)
        elif p.startswith('<h3>'):
            doc.add_heading(p[4:-5], level=3)
        elif p.startswith('<h4>'):
            doc.add_heading(p[4:-5], level=4)
        elif p.startswith('<ul>'):
            continue
        elif p.startswith('<li>'):
            doc.add_paragraph(p[4:-5], style='List Bullet')
        else:
            if p.strip():
                doc.add_paragraph(p)
    
    # 워드 파일 저장
    doc.save(docx_file)

if __name__ == '__main__':
    current_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(current_dir, '품질관리시스템_기능설명서.md')
    docx_file = os.path.join(current_dir, '품질관리시스템_기능설명서.docx')
    
    convert_md_to_docx(md_file, docx_file)
    print(f"변환 완료: {docx_file}") 